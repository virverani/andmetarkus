# Install necessary libraries

import requests
import json
import pandas as pd
import matplotlib.pyplot as plt

# url = "https://demo-datahub.rik.ee/api/v1/meta/classifications"

url_riik = 'https://api.worldbank.org/v2/countries/EST/?format=json'
url_rahvastik_ = 'https://api.worldbank.org/v2/country/EST/indicator/SP.POP.TOTL?format=json'
url_women = "https://api.worldbank.org/v2/country/EST/indicator/SP.POP.TOTL.FE.IN?format=json"

response = requests.get(url_rahvastik_)
data = response.json()

response_woman = requests.get(url_women)
data_woman = response_woman.json()


# json dumps muudab väljundi terminalis loetavaks
# print(json.dumps(data, indent=2, ensure_ascii=False))

values = {'year': [], 'population': [], 'women_population': []}

# {"year": ["2021", "2020", "2019", ...], "population": [1331057, 1326535, 1324820, ...]}

for item in data[1]:
    values['year'].append(item['date'])
    values['population'].append(item['value'])

for women in data_woman[1]:
    values['women_population'].append(women['value'])

# print(json.dumps(values, indent=2, ensure_ascii=False))

df = pd.DataFrame(values)

df = df.sort_values(by='year')  # sorteerin aasta järgi kasvavalt

df['men_population'] = df['population'] - df['women_population']

# väljastan esimesed read
print(df.head())

# joonistamise osa
df.plot(x='year', y=['population', 'women_population', 'men_population'], kind='line',
        title='Estonian population', xlabel='Year', ylabel='Population')
plt.ylim(bottom=0)  # y telje algus


plt.savefig("EE_population_men_women.png")
plt.show()


# Save to word file
from docx import Document

# Create a new Word document
doc = Document()
doc.add_heading('Estonian population', level=1)
p = doc.add_paragraph(
    'You can see the Estonian total population and population by gender.')

# Lisa graafik Wordi dokumendile
doc.add_picture('EE_population_men_women.png')


# add page break
doc.add_page_break()

doc.add_heading('Data Sources:', level=2)
doc.add_paragraph(
    'Source: The World Bank (https://data.worldbank.org/indicator/SP.POP.TOTL?locations=EE)')
doc.add_paragraph(
    'Source: The World Bank (https://data.worldbank.org/indicator/SP.POP.TOTL.FE.IN?locations=EE)')

# Add a table to the document
table = doc.add_table(rows=1, cols=len(df.columns))
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
for i, col_name in enumerate(df.columns):
    hdr_cells[i].text = str(col_name)

for _, row in df.iterrows():
    row_cells = table.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = str(value)


# Save the document
doc.save("EE_population_analysis.docx")

